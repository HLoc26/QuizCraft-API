import io
import os
import docx
import pymupdf
from ..constants import UPLOAD_FOLDER
from .ocr_service import OCRService
from PIL import Image


class ScanService:
    def __init__(self):
        self.UPLOAD_FOLDER: str = UPLOAD_FOLDER
        self.ocr_service = OCRService()

    def get_path(self, filename: str) -> str:
        return os.path.join(self.UPLOAD_FOLDER, filename)

    async def scan_pdf_pages(self, file_or_path):
        results = []
        if isinstance(file_or_path, str):
            doc = pymupdf.open(self.get_path(file_or_path))
        else:
            file_or_path.seek(0)
            doc = pymupdf.open(stream=file_or_path.read(), filetype="pdf")

        for i, page in enumerate(doc, start=1):
            extracted = page.get_text("text")
            if extracted.strip():
                text = extracted
            else:
                pix = page.get_pixmap(dpi=300)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                text = await self.ocr_service.use_easyocr(img)
            results.append({"page": i, "text": text})
        return results

    async def scan_pdf(self, file_or_path):
        text: str = ""

        if isinstance(file_or_path, str):  # path
            doc = pymupdf.open(self.get_path(file_or_path))
        else:  # file-like object
            # Đảm bảo seek về đầu
            file_or_path.seek(0)
            doc = pymupdf.open(stream=file_or_path.read(), filetype="pdf")

        for page in doc:
            extracted = page.get_text("text")
            if extracted.strip():
                text += extracted + "\n"
            else:
                pix = page.get_pixmap(dpi=300)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_text = await self.ocr_service.use_easyocr(img)
                text += ocr_text + "\n"
        return text

    async def scan_image(self, file_or_path):
        if isinstance(file_or_path, str):
            file_path = self.get_path(file_or_path)
            return await self.ocr_service.use_easyocr(file_path)
        else:
            file_or_path.seek(0)
            img = Image.open(file_or_path)
            return await self.ocr_service.use_easyocr(img)

    async def scan_text(self, file_or_path):
        if isinstance(file_or_path, str):
            file_path = self.get_path(file_or_path)
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        else:
            file_or_path.seek(0)
            return file_or_path.read().decode("utf-8")

    async def scan_docx(self, file_or_path):
        if isinstance(file_or_path, str):
            file_path = self.get_path(file_or_path)
            doc = docx.Document(file_path)
        else:
            file_or_path.seek(0)
            doc = docx.Document(file_or_path)

        return "\n".join([para.text for para in doc.paragraphs])
